import os

import fitz
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from core.ocr_engine import configure_tesseract


def export_txt(path, text):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)


def export_docx(path, text):
    doc = Document()
    for line in text.splitlines():
        paragraph = doc.add_paragraph(line)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    if not text.strip():
        doc.add_paragraph("")
    doc.save(path)


def _pdf_font_name(page):
    fontfile = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts", "tahoma.ttf")
    if os.path.isfile(fontfile):
        page.insert_font(fontname="Tahoma", fontfile=fontfile)
        return "Tahoma"
    return "helv"


def export_text_pdf(path, text):
    doc = fitz.open()
    page_width, page_height = fitz.paper_size("a4")
    margin = 50
    fontsize = 12
    line_height = fontsize * 1.5
    lines = (text or "").splitlines() or [""]

    page = doc.new_page(width=page_width, height=page_height)
    fontname = _pdf_font_name(page)
    y = margin

    for line in lines:
        if y + line_height > page_height - margin:
            page = doc.new_page(width=page_width, height=page_height)
            fontname = _pdf_font_name(page)
            y = margin
        page.insert_text(
            (page_width - margin, y),
            line,
            fontsize=fontsize,
            fontname=fontname,
        )
        y += line_height

    doc.save(path)


def export_searchable_image_pdf(path, image, tesseract_path, lang, tesseract_config):
    pytesseract = configure_tesseract(tesseract_path)
    pdf_bytes = pytesseract.image_to_pdf_or_hocr(
        image,
        lang=lang,
        config=tesseract_config,
        extension="pdf",
    )
    with open(path, "wb") as f:
        f.write(pdf_bytes)


def save_export(path, text, fmt, source_image=None, tesseract_path=None, lang="fas+eng", tesseract_config=None):
    fmt = (fmt or "txt").lower().lstrip(".")
    if fmt == "txt":
        export_txt(path, text)
        return "txt"
    if fmt == "docx":
        export_docx(path, text)
        return "docx"
    if fmt == "pdf":
        if source_image is not None and tesseract_path and tesseract_config:
            try:
                export_searchable_image_pdf(
                    path, source_image, tesseract_path, lang, tesseract_config,
                )
                return "pdf_searchable"
            except Exception:
                pass
        export_text_pdf(path, text)
        return "pdf_text"
    raise ValueError(f"Unsupported export format: {fmt}")
